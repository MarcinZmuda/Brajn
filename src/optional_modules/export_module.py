"""
Export Module — optional module.
Export articles to DOCX, HTML, TXT formats.
"""
import re
import io
from fastapi.responses import StreamingResponse


def _markdown_to_html(text: str) -> str:
    """Simple markdown to HTML conversion."""
    html = text
    html = re.sub(r"^# (.+)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)
    html = re.sub(r"^## (.+)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
    html = re.sub(r"^### (.+)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
    html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
    html = re.sub(r"\*(.+?)\*", r"<em>\1</em>", html)

    paragraphs = html.split("\n\n")
    result = []
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if p.startswith("<h") or p.startswith("<ul") or p.startswith("<ol"):
            result.append(p)
        else:
            result.append(f"<p>{p}</p>")

    return "\n".join(result)


def export_docx(text: str) -> StreamingResponse:
    """Export article as DOCX file."""
    try:
        from docx import Document

        doc = Document()

        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                # Clean markdown bold/italic
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                doc.add_paragraph(clean)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=article.docx"},
        )

    except ImportError:
        return {"error": "python-docx not installed"}


def export_html(text: str) -> dict:
    """Export article as HTML."""
    html = _markdown_to_html(text)
    full_html = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1 {{ font-size: 2em; margin-bottom: 0.5em; }}
        h2 {{ font-size: 1.5em; margin-top: 1.5em; }}
        h3 {{ font-size: 1.2em; }}
        p {{ margin-bottom: 1em; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""
    return {"html": full_html}


def export_txt(text: str) -> dict:
    """Export article as plain text."""
    clean = re.sub(r"[#*_]", "", text)
    return {"text": clean}
